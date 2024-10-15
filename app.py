import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import streamlit as st
import html
import re  # Import regex for better matching

# Function to extract content from <h1> to <h6>, <p> tags and <a> links
def extract_content_from_url(url):
    response = requests.get(url)
    response.encoding = 'utf-8'
    soup = BeautifulSoup(response.text, 'html.parser')

    content = []
    h1_text = ""
    start = False
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
        if element.name == 'h1':
            h1_text = element.get_text().strip()  # Capture the H1 text
            start = True
        if start:
            text = element.get_text().strip()
            text = html.unescape(text)  # Converts HTML entities into normal characters
            if text:
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    content.append({'type': 'heading', 'level': element.name, 'text': text})
                elif element.name == 'p':
                    paragraph = ""
                    for sub_element in element:
                        if sub_element.name == 'a' and sub_element.get('href'):
                            # Create hyperlink text in the format: 'Text (URL)'
                            anchor_text = sub_element.get_text()
                            url = sub_element.get("href")
                            paragraph += f'{anchor_text} ({url}) '  # Include link as text
                        else:
                            paragraph += sub_element.string if sub_element.string else ''
                    content.append({'type': 'paragraph', 'text': paragraph.strip()})

    return content, h1_text  # Return the H1 text as well

# Function to create a Word document from the extracted content
def create_word_file(file_name, content, url):
    doc = Document()
    
    # Add the URL as the first paragraph
    doc.add_paragraph(f"URL: {url}")

    # Add extracted content to the Word file with formatting
    for element in content:
        if element['type'] == 'heading':
            level = int(element['level'][1])  # Heading level (h1 = 1, h2 = 2, etc.)
            doc.add_heading(element['text'], level=level)
        elif element['type'] == 'paragraph':
            paragraph = doc.add_paragraph()
            # Regex pattern to capture anchor text and URL
            pattern = re.compile(r"(.*?)(\s+\((https?://[^\s]+)\))?$")
            parts = element['text'].split('. ')
            for part in parts:
                match = pattern.match(part.strip())
                if match:
                    anchor_text = match.group(1).strip()  # Get anchor text
                    url = match.group(3) if match.group(3) else ''  # Get URL if available
                    # Add hyperlink to the document only if the URL is present
                    if url:
                        add_hyperlink(paragraph, anchor_text, url)
                    else:
                        paragraph.add_run(anchor_text + ' ')  # Add normal text if no link
                else:
                    paragraph.add_run(part + ' ')  # Add normal text if no match

    # Save the Word file
    doc.save(file_name)
    return file_name

# Function to add a hyperlink to a paragraph
def add_hyperlink(paragraph, text, url):
    # This function adds a clickable hyperlink to a paragraph
    part = paragraph.part
    r_id = part.relate_to(url, docx.oxml.ns.RT.HYPERLINK, is_external=True)
    
    # Create the w:hyperlink tag and add needed attributes
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create the w:r element (run), where we put the text of the hyperlink
    run = OxmlElement('w:r')
    
    # Add the text to the run
    rPr = OxmlElement('w:rPr')  # Add hyperlink style (This will make it blue and underlined)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    # Add the run to the hyperlink element, and then to the paragraph
    hyperlink.append(run)
    paragraph._element.append(hyperlink)

# Streamlit interface
st.title("HTML Content Extractor to Word")
url = st.text_input("Enter the page URL")
jira_link = st.text_input("Add the JIRA link (TT - Traffic Team)")

if st.button("Generate Word File"):
    if url:
        # Extract content from the URL
        content, h1_text = extract_content_from_url(url)
        if content:
            # Determine the file name based on the JIRA link
            if jira_link:
                ticket_number = jira_link[-4:]
                filename = f"Brief SEO Optimization - TT-{ticket_number}.docx"
            else:
                filename = f"{h1_text}.docx"  # Use H1 text if JIRA link is empty

            # Create the Word file
            create_word_file(filename, content, url)  # Pass URL to the function

            with open(filename, "rb") as file:
                # Automatically trigger the download
                st.download_button(
                    label="Your file is ready! Click to download",
                    data=file,
                    file_name=filename
                )
        else:
            st.error("Unable to extract content from this URL.")
    else:
        st.error("Please fill out the URL field.")
